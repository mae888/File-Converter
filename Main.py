import PyPDF2
from docx import Document
from docx.shared import Inches
import sys
import os
import comtypes.client
from docx2pdf import convert

def removeQuotes (s):
    s = s.replace("\"", '')
    s = s.replace("\"", '')
    return s

def editPath (path):
    path = removeQuotes(path)
    text = ""
    for char in path:
        if char == "\\":
            text += "\\\\"
        else :
            text += char
    return text

def mergePDF ():
    numOfFiles = int(input("Enter the number of files: "))
    files = []
    for i in range(numOfFiles):
        path = input("Enter the path to file " + str(i+1) +": ")
        path = editPath(path)
        files.append(path)
    merger = PyPDF2.PdfMerger()
    for pdf in files:
        merger.append(open(pdf, 'rb'))
    output = open("Merged Files.pdf", 'wb')
    merger.write(output)
    output.close()

def mergeWord():
    numOfFiles = int(input("Enter the number of files: "))
    files = []
    for i in range(numOfFiles):
        path = input("Enter the path to file " + str(i + 1) + ": ")
        path = editPath(path)
        files.append(path)
    merged_document = Document()
    for file in files:
        document = Document(file)
        for paragraph in document.paragraphs:
            merged_document.add_paragraph(paragraph.text)
        for table in document.tables:
            merged_document.add_table(table)
    merged_document.save('Merged Word Documents.docx')

def convW2PDF():
    print("**** All the DOCX files must be in a single folder ****")
    folderPath = input("Enter the path to the folder: ")
    folderPath = editPath(folderPath)
    convert(folderPath, "C:\\Users\\acer\\Desktop\\")

def convPP2PDF():
    print("**** All the PPTX files must be in a single folder ****")
    folderPath = input("Enter the path to the folder: ")
    folderPath = editPath(folderPath)
    input_folder_path = folderPath
    output_folder_path = "C:\\Users\\acer\\Desktop\\"
    #input_folder_path = os.path.abspath(input_folder_path)
    output_folder_path = os.path.abspath(output_folder_path)

    input_file_paths = os.listdir(input_folder_path)
    for input_file_name in input_file_paths:
        if not input_file_name.lower().endswith((".ppt", ".pptx")):
            continue
        input_file_path = os.path.join(input_folder_path, input_file_name)
        powerpoint = comtypes.client.CreateObject("Powerpoint.Application")
        powerpoint.Visible = 1
        slides = powerpoint.Presentations.Open(input_file_path)
        file_name = os.path.splitext(input_file_name)[0]
        output_file_path = os.path.join(output_folder_path, file_name + ".pdf")
        slides.SaveAs(output_file_path, 32)
        slides.Close()


while True:
    print("Choose a Number:")
    print("[1] Merge PDF Files")
    print("[2] Merge Word Documents")
    print("[3] Convert Word Documents to PDF Files")
    print("[4] Convert PowerPoint Documents to PDF Files")
    print("[5] Exit")
    c = int(input("Choice: "))
    if c == 1:
        mergePDF()
        print("Success!\n")
    elif c == 2:
        mergeWord()
        print("Success!\n")
    elif c == 3:
        convW2PDF()
        print("Success!\n")
    elif c == 4:
        convPP2PDF()
        print("Success!\n")
    elif c == 5:
        print("\nExiting...\n")
        break
    else:
        print("Incorrect Input!\n")




